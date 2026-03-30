from __future__ import annotations

import csv
import re
from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx"}


def sanitize_filename(filename: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", filename.strip())
    return cleaned or "upload"


def load_documents(file_path: Path) -> list[Document]:
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {supported}")

    if suffix == ".pdf":
        return _load_pdf(file_path)
    if suffix == ".docx":
        return _load_docx(file_path)
    if suffix in {".txt", ".md"}:
        return _load_text(file_path)
    if suffix == ".csv":
        return _load_csv(file_path)
    if suffix == ".xlsx":
        return _load_xlsx(file_path)
    raise ValueError(f"Unsupported file type '{suffix}'")


def _load_pdf(file_path: Path) -> list[Document]:
    reader = PdfReader(str(file_path))
    documents: list[Document] = []

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={"page": index, "source_type": "pdf"},
            )
        )
    return documents


def _load_docx(file_path: Path) -> list[Document]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ValueError(
            "DOCX support requires 'python-docx'. Install backend dependencies first."
        ) from exc

    document = DocxDocument(str(file_path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs).strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source_type": "docx"})]


def _load_text(file_path: Path) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source_type": "text"})]


def _load_csv(file_path: Path) -> list[Document]:
    rows: list[str] = []
    with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for index, row in enumerate(reader, start=1):
            row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
            if row_text:
                rows.append(f"Row {index}: {row_text}")

    if not rows:
        return []

    return [
        Document(
            page_content="\n".join(rows),
            metadata={"source_type": "csv", "row_range": f"1-{len(rows)}"},
        )
    ]


def _load_xlsx(file_path: Path) -> list[Document]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError(
            "Excel support requires 'openpyxl'. Install backend dependencies first."
        ) from exc

    workbook = load_workbook(filename=str(file_path), read_only=True, data_only=True)
    documents: list[Document] = []

    for sheet in workbook.worksheets:
        rows: list[str] = []
        row_count = 0
        for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(cell).strip() for cell in row if cell not in (None, "")]
            if not values:
                continue
            row_count += 1
            rows.append(f"Row {index}: {' | '.join(values)}")

        if rows:
            documents.append(
                Document(
                    page_content="\n".join(rows),
                    metadata={
                        "source_type": "xlsx",
                        "sheet": sheet.title,
                        "row_range": f"1-{row_count}",
                    },
                )
            )

    return documents
